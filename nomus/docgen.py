"""Генерация заявления в госинспекцию труда: текст + .docx (FR-30..FR-33).

Поля заявления держатся только в оперативной памяти FSM и не пишутся в БД
(требование приватности §7 ТЗ).
"""

import io
from dataclasses import dataclass, field
from datetime import date


@dataclass
class ClaimFields:
    full_name: str = ""
    citizenship: str = ""
    employer: str = ""
    work_period: str = ""
    debt_amount: str = ""
    contact: str = ""
    violated_articles: list[str] = field(default_factory=list)  # FR-33: из ответа RAG


CLAIM_TEMPLATE = """В Государственную инспекцию труда
{region}

от {full_name},
гражданина(ки) {citizenship},
контакт: {contact}

ЗАЯВЛЕНИЕ
о невыплате заработной платы

Я, {full_name}, работал(а) у работодателя «{employer}» в период {work_period}.

Работодатель не выплатил мне заработную плату. Сумма задолженности составляет
{debt_amount} тенге.
{articles_block}
На основании изложенного ПРОШУ:
1. Провести проверку соблюдения трудового законодательства работодателем
   «{employer}».
2. Принять меры по восстановлению моих трудовых прав и взысканию
   задолженности по заработной плате.

Приложения: копии имеющихся документов (при наличии).

Дата: {today}
Подпись: ______________ / {full_name} /
"""


def build_claim_text(f: ClaimFields, region: str = "по месту нахождения работодателя") -> str:
    if f.violated_articles:
        articles = "; ".join(f.violated_articles)
        articles_block = (
            f"\nСчитаю, что работодателем нарушены требования законодательства "
            f"Республики Казахстан: {articles}.\n"
        )
    else:
        articles_block = "\n"
    return CLAIM_TEMPLATE.format(
        region=region,
        full_name=f.full_name or "____________________",
        citizenship=f.citizenship or "____________",
        employer=f.employer or "____________",
        work_period=f.work_period or "____________",
        debt_amount=f.debt_amount or "________",
        contact=f.contact or "____________",
        articles_block=articles_block,
        today=date.today().strftime("%d.%m.%Y"),
    )


def build_claim_docx(f: ClaimFields) -> bytes:
    """Возвращает .docx как bytes (для отправки BufferedInputFile в aiogram)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for para in build_claim_text(f).split("\n\n"):
        p = doc.add_paragraph()
        p.add_run(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
